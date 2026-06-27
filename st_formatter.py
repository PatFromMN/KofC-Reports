import pandas as pd
from docx import Document

FORM_DUE_DATES = {
    "SP-7": "June 30",
    "185": "June 30",
    "365": "June 30",
    "1728": "January 31",
}

REPORT_SUMMARIES = {
    'form description' : "Each Council is required to submit the following forms by the due date to be eligible for Star Council. The forms are: Report of Council Officers (#185) which is due by June 30, Service Program Personnel report (#365) which is due by June 30, and the Annual Survey of Fraternal Activity (#1728) which is due by January 31. The SP-7 forms are due by June 30 are due by the end of the fraternal year.",
    'star council description' : "Add some text to describe the Star Council requirements.",
    'mcgivny description': "Tp earn this award, a council must meet or exceed its yearly memership quota.",
    'columbian description' : "To earn this award, a council must conduct and report programs in each of the rour program categories: Faith, Family, Community, and Life.",
    'founders description' : "The Founders Award recognizes excellence in the promotion of our insurance and fraternal benefits."
}
tracker_name = ''

def get_file():
    """ Download the Star Tracker field identified by the user."""
    global tracker_name
    
    tracker_name = input("Enter the Tracker file name and extension: ")
    
    try:
        df = pd.read_excel(tracker_name)
        return df
    except FileNotFoundError:
        print(f"Could not find the file '{tracker_name}'.")
    except ValueError:
        print(f"'{tracker_name}' does not look like a valid Excel file. Make sure thextension and format are correct.")
    except PermissionError:
        print(f"Permission denied when trying to open '{tracker_name}'. Is it open in another program?")
    except Exception as e:
        print(f"An unexpected error occured: {e}")
        
        
def get_district_info():
    """The routine extractrs the information for the district identified by the user."""
    district_number = input("Enter the District number: ")
    district_id = int(input("Enter the District number from Star Tracker: "))
    return district_number, district_id

def get_district_data(st_df, district_id):
    """This function extracts the data from the dataframe and returns the data for the district in question."""
    
    district_data = {}
    
    for _, record in st_df.iterrows():
        if record['District #'] == district_id:
            key = record['Council #']
            district_data[key] = record.to_dict()
            # print(record)
    
    return district_data

def _calculate_membership_status(district_data):
    """This function calculates the membership quota and gain for the district."""
    district_quota = 0
    district_gain = 0
    for council_number, record in district_data.items():
        district_quota += record['Membership Quota']
        district_gain += record['Membership Gain (YTD)']

    return district_quota, district_gain

def add_membership_status(document, district_data):
    """This function adds the membership data to the report."""
    document.add_heading(f"District Membership Status", level=1)
    district_quota, district_gain =_calculate_membership_status(district_data)
    document.add_paragraph(f"District Membership Quota: {district_quota:.0f}")
    document.add_paragraph(f"District Membership Gain: {district_gain:.0f}")
    document.add_paragraph(f"District Membership Quota attained: {district_gain/district_quota:.2%}")

def add_form_status(document, district_data):
    """This function adds the form status to the report."""
    # Set up this section of the report.
    document.add_heading(f"District Form Status", level=1)
    section_description = REPORT_SUMMARIES["form description"]
    document.add_paragraph(section_description)

    # Set up the form tracking table
    rows = 5
    columns = len(district_data) + 2                        # Plus 2 for the form titles and the due dates.
    table = document.add_table(rows=rows, cols=columns)
    table.style = 'Table Grid'
    table.cell(1,0).text = "SP-7"
    table.cell(1,columns-1).text = FORM_DUE_DATES['SP-7']
    table.cell(2,0).text = "#185"
    table.cell(2, columns-1).text = FORM_DUE_DATES['185']
    table.cell(3,0).text = "#365"
    table.cell(3, columns-1).text = FORM_DUE_DATES['365']
    table.cell(4,0).text = "#1728"
    table.cell(4, columns-1).text = FORM_DUE_DATES['1728']
    table.cell(0,columns-1).text = "Due Date"
    
    # Add the Council numbers to the first row and the data in the appropriate cells.
    for i, council_number in enumerate(district_data.keys()):
        table.cell(0, i+1).text = f"{district_data[council_number]['Council #']}"
        table.cell(1, i+1).text = f"{district_data[council_number]['Form SP-7 Rec\'d']}"
        table.cell(2, i+1).text = f"{district_data[council_number]['Form 185 (Officers Chosen)']}"
        table.cell(3, i+1).text = f"{district_data[council_number]['Form 365 (Program Personnel List)']}"
        table.cell(4, i+1).text = f"{district_data[council_number]['Form 1728 \n(Annual Survey)']}"

def print_safe_environment_status(document, district_data):
    """This function adds the Safe Environment status to the report."""
    
    document.add_heading(f"Safe Environment Status", level=1)
    for council_number, record in district_data.items():
        if record['OYP Status'] != "Compliant":
            safe_environment_status = "Incomplete"
        else:
            safe_environment_status = "Compliant"
        document.add_paragraph(f"Council {district_data[council_number]['Council #']} is {safe_environment_status}.")

def print_star_council_status(document, district_data):
    """This function adds the Star Council status to the report."""
    document.add_heading(f"Star Council Award Status", level=1)

    # Add the table to show the Star Council status
    rows = 5
    columns = len(district_data) + 1                        # Plus one column for the award titles
    table = document.add_table(rows=rows, cols=columns)
    table.style = 'Table Grid'
    table.cell(1,0).text = "McGivney Award"
    table.cell(2,0).text = "Columbian Award"
    table.cell(3,0).text = "Founders Award"
    table.cell(4,0).text = "Star Council?"

    # Add the Council information
    for i, council_number in enumerate(district_data.keys()):
        table.cell(0, i+1).text = f"{district_data[council_number]['Council #']}"
        if district_data[council_number]['McGivney Award?'] == "YES":
            table.cell(1, i+1).text = "Yes"
        else:
            table.cell(1, i+1).text = "No"

        if district_data[council_number]['Columbian Award?'] == "YES":
            table.cell(2, i+1).text = "Yes"
        else:
            table.cell(2, i+1).text = "No"
        
        if district_data[council_number]['Founders Award?'] == "YES":
            table.cell(3, i+1).text = "Yes"
        else:
            table.cell(3, i+1).text = "No"
        
        if district_data[council_number]['Star Council?'] == "YES":
            table.cell(4, i+1).text = "Yes"
        else:
            table.cell(4, i+1).text = "No"
    

def print_district_report(district_data, district_number):
    """Controls the printing of the District Deputy's report for the District."""
    
    # Get the file name from the user. 
    file_name = input("Enter a file name for the report (ex: District Report dd-mm-yyyy): ")
    document = Document()

    document.add_heading(f"District {district_number} Report", 0)

    # generate the District Deputy's report with the data from the Star Tracker file. 
    add_membership_status(document, district_data)
    add_form_status(document, district_data)
    print_safe_environment_status(document, district_data)
    print_star_council_status(document, district_data)


    document.add_heading(f"Field Agent Contact Information", level=1)
    for council_number, record in district_data.items():
        document.add_heading(f"Council {council_number}", level=2)
        document.add_paragraph(f"Field Agent: {record['Council Field Agent Name']},  email:{record['Council Field Agent Email']}")

    # save the completed report
    document.save(file_name + ".docx")

def main():
    # Print the welcome message to the command line.
    print("Welcome to the Star Tracker formatter.")
    print("You can download the lastest Start Tracker file from: https://www.californiaknights.org/2025-2026-council-star-awards/")
    st_df = get_file()
    district_number, district_id = get_district_info()

    print(f"Processing the data from District {district_number}")
    district_data = get_district_data(st_df, district_id)
    # print(district_data)

    print_district_report(district_data, district_number)

if __name__ == "__main__":
    main()    