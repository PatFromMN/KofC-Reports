# This program extracts a council's information from the Star Tracker spreadsheet.
# The data is used to create a Word document that can be shared with the Council.


import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime as dt

FORM_DUE_DATES = {
    "SP-7": "June 30",
    "185": "June 30",
    "365": "June 30",
    "1728": "January 31",
}

tracker_name = ''

def get_file():
    """ Download the Star Tracker field identified by the user."""
    global tracker_name
    
    tracker_name = input("Enter the Tracker file name and extension: ")
   
    try:
        df = pd.read_excel(tracker_name)
        return df
    except FileNotFoundError:
        print(f"Could not find the file '{tracker_name}'.")
    except ValueError:
        print(f"'{tracker_name}' does not look like a valid Excel file. Make sure thextension and format are correct.")
    except PermissionError:
        print(f"Permission denied when trying to open '{tracker_name}'. Is it open in another program?")
    except Exception as e:
        print(f"An unexpected error occured: {e}")
        
def get_council_info():
    """This function gets the council number from the user, then returns the number to the calling routine."""

    council_id = int(input("Enter the Council id number: "))
    return council_id

def get_council_data(star_tracker_df, council_id):
    """This function extracts the data from the dataframe and returns the council's data."""

    if star_tracker_df is None:
        print("Error: no data loaded.")
        council_data = {}
    else:
        council_data = {}
        for _, record in star_tracker_df.iterrows():
            if record['Council #'] == council_id:
                council_data = record.to_dict()

    return council_data

def print_council_report(st_df, council_id):
    """This function prints the Council Report"""

    # set up the document
    document = Document()
    todays_date = dt.now()

    file_name = f"Council {council_id} Report of {todays_date: %B %d, %Y}.docx"
    
    # print the report header
    document.add_heading(f"Council {council_id} Status Report", level=0)
    document.add_heading(f"Generated on: {todays_date: %d %B %Y}", level=1)
    # print the annual report status
    document.add_heading("Council Forms", level=2)
    if st_df['Form 185 (Officers Chosen)'] == 'YES':
        document.add_paragraph("Report of Council Officers (#185) was submitted.", style='List Bullet')
    else:
        document.add_paragraph("Report of Council Officers (#185) was not submitted.", style='List Bullet')
    if st_df['Form 365 (Program Personnel List)'] == 'YES':
        document.add_paragraph("Service Program Personnel Report (#365) was submitted.", style='List Bullet')
    else:
        document.add_paragraph("Service Program Personnel Report (#365) was not submitted.", style='List Bullet')
    if st_df['Form 1728 \n(Annual Survey)'] == 'YES':
        document.add_paragraph(f"Annual Survey of Fraternal Activity (#1728) was submitted.", style='List Bullet')
    else:
        document.add_paragraph(f"Annual Survey of Fraternal Activity (#1728) was not submitted.", style='List Bullet')
    if st_df["Form SP-7 Rec'd"] == 'APPROVED':
            document.add_paragraph(f"Columbian Award Application (#SP-7) was approved.", style='List Bullet')
    else:
        document.add_paragraph(f"Columbian Award Application (#SP-7) was not approved.", style='List Bullet')

    # print the Fraternal Benefits
    document.add_heading("Fraternal Benefits", level=2)
    # p = document.add_paragraph(f"The Council {council_id} Field Agent is {st_df['Council Field Agent Name']} and can be reached at {st_df['Council Field Agent Email']}.")
    # p.add_run(f"and {st_df['Total FBE']:.0f} events. They had a total {st_df['Total FBE Attendees']:.0f} attendees since the beginning of the year.")
    
    p1 = document.add_paragraph(f"The Council {council_id} Field Agent is {st_df['Council Field Agent Name']} and can be reached at {st_df['Council Field Agent Email']}. ")
    p1.add_run(f"Council {council_id} has a Fraternal Benefits quota of {st_df['FBE Attendee Quota']:.0f} members. ")
    p1.add_run(f"They held {st_df['Total FBE']:.0f} events. ")
    p1.add_run(f"They had a total {st_df['Total FBE Attendees']:.0f} attendees since the beginning of the year.")

    # print the Star Council Status
    document.add_heading("Star Council", level=2)
    if st_df['OYP Status'] == 'Compliant':
        document.add_paragraph(f"Council {council_id} complies with the Safe Environment requirements.", style='List Bullet')
    else:
        document.add_paragraph(f"Council {council_id} does not comply with the Safe Environment requirements.", style='List Bullet')
    if st_df['McGivney Award?'] == "YES":
        document.add_paragraph(f"Council {council_id} received the McGivney Award.", style='List Bullet')
    else:
        document.add_paragraph(f"Council {council_id} did not receive the McGivney Award.", style='List Bullet')
    if st_df['Columbian Award?'] == 'YES':
        document.add_paragraph(f"Council {council_id} received the Columbian Award.", style='List Bullet')
    else:
        document.add_paragraph(f"Council {council_id} did not receive the Columbian Award.", style='List Bullet')
    if st_df['Founders Award?'] == 'YES':
        document.add_paragraph(f"Council {council_id} received the Founders Award.", style='List Bullet')
    else:
        document.add_paragraph(f"Council {council_id} did not receive the Founders Award.", style='List Bullet')
    if st_df['Star Council?'] == "YES":
        document.add_paragraph(f"Council {council_id} received the Star Council Award.", style='List Bullet')
    else:
        document.add_paragraph(f"Council {council_id} did not receive the Star Council Award.", style='List Bullet')

    # save the completed report
    document.save(file_name)

def main():
    print("Welcome to the Council Tracker program.")
    print("You can download the lastest Start Tracker file from: https://www.californiaknights.org/2025-2026-council-star-awards/")
    star_tracker_df = get_file()
    council_id = get_council_info()
    council_data = get_council_data(star_tracker_df, council_id)
    
    print_council_report(st_df=council_data, council_id=council_id)

if __name__ == "__main__":
    main()